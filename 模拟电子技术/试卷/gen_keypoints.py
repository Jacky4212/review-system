#!/usr/bin/env python3
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

doc = Document()
s = doc.styles['Normal']
s.font.name = 'Microsoft YaHei'
s.font.size = Pt(11)
doc.add_heading('模拟电子技术 - 关键注意事项', level=0).alignment = WD_ALIGN_PARAGRAPH.CENTER
doc.add_paragraph('按章节整理易错点、必记公式、关键区别', style='Intense Quote')

data = {
'第1章 绪论': [
('易错','"注意：输入、输出电阻为交流电阻"（ch01 1.5）--不是直流电阻！'),
('区别','频率失真(线性失真)不产生新频率分量 vs 非线性失真产生新频率分量(谐波)。'),
('公式','电压增益dB=20lg|Av|；电流dB=20lg|Ai|；功率dB=10lgAp。电压电流20倍，功率10倍。'),
('必记','四种放大模型理想条件：电压(Ri->inf,Ro->0)；电流(Ri->0,Ro->inf)。'),
],
'第2章 运算放大器': [
('必记','"虚短和虚断是用来分析各种运放线性应用电路的有力法则，必须熟练掌握。"虚短!=真短路，虚断!=真断路。'),
('区别','同相Av=1+R2/R1(正),Ri->inf,有共模电压；反相Av=-Rf/R1(负),Ri=R1,虚地无共模电压。'),
('易错','电压跟随器(Av=1)作用是"隔离或缓冲"，不是放大电压。'),
('公式','积分：vo=-(1/RC)积分vi dt；微分：vo=-RC dvi/dt。'),
],
'第3章 二极管': [
('区别','P型多子=空穴；N型多子=自由电子。'),
('区别','雪崩(少子获动能撞击)vs齐纳(强电场破坏共价键)。电击穿可逆，热击穿不可逆。'),
('易错','小信号模型前提：正向偏置且vD>>VT(=26mV)。rd=VT/IDQ，与Q点有关。'),
('易错','齐纳二极管工作在反向电击穿状态，不是正向导通！'),
],
'第4章 MOS场效应管': [
('区别','增强型"必须依靠栅极外加电压才能产生反型层"VGS=0无沟道；耗尽型VGS=0已有沟道。'),
('必记','MOSFET单极型器件--沟道仅一种载流子。VCCS电压控制电流源。'),
('易错','恒流区条件vGS>VTN且vDS>=vGS-VTN。"必须检验是否满足"如不满足"必须作出新的假设"。'),
('公式','gm=2Kn(VGSQ-VTN)=2sqrt(Kn IDQ)。'),
('易错','VDS极性接反->PN结正偏->MOS管失效！NMOS漏极为正。'),
('必记','共源(反相Av=-gmRd)；共漏(同相Av=1)；共栅(同相)。'),
],
'第5章 BJT': [
('必记','BJT放大条件：发射结正偏+集电结反偏。NPN(VC>VB>VE)，PNP全部相反。'),
('区别','BJT(CCCS,beta ib)vs FET(VCCS,gm vgs)。BJT双极Ri中等，FET单极Ri极高。'),
('公式','rbe=200+(1+beta)26mV/IEQ。不能万用表量b/e电阻！'),
('必记','共射(反相高Av)；共集/射随(同相Av=1缓冲)；共基(同相低Ri)。'),
],
'第7章 模拟集成电路': [
('必记','"利用叠加原理将差模和共模分开处理--差分放大分析的核心手段。"差模时源极虚地。'),
('区别','差模vid=vP-vN(大小相等相位相反)；共模vic=(vP+vN)/2(大小相等相位相同)。'),
('公式','Avd=-gmRd；KCMR=|Avd/Avc|越大越好。'),
('易错','温度变化+电源波动=共模信号->差分放大可抑制=克服零漂的原理。'),
],
'第8章 反馈放大电路': [
('必记','瞬时极性法四步：区分电路和反馈网络->沿信号方向逐级标注->知道相位关系->正确确定反馈引入位置。'),
('区别','串联->Ri增大；并联->Ri减小；电压->Ro减小；电流->Ro增大。'),
('公式','Af=A/(1+AF)；深度负反馈(AF>>1)：Af=1/F。'),
('易错','深度负反馈近似有误差！环路增益越大误差越小。高频时A下降AF减小。'),
('必记','自激振荡AF=-1(|AF|=1+相移180度)。信源=噪声+干扰。解决=密勒补偿。'),
],
'第9章 功率放大电路': [
('区别','甲类(360度)效率~25%；乙类(180度)交越失真 eta_max=78.5%；甲乙类无交越；丁类>90%。'),
('公式','OCL:Pomax=VCC^2/(2RL)；OTL:Pomax=VCC^2/(8RL)。'),
('必记','交越失真->甲乙类偏置(二极管偏置/VBE倍增电路)。'),
('易错','OCL输出级共集电极Av=1，无电压放大功能！'),
],
'第10章 信号处理与产生': [
('区别','LPF低通/HPF高通/BPF带通(串联)/BRF带阻(并联)/APF全通。'),
('必记','Q=0.707巴特沃斯响应通带最平坦无过冲。'),
('公式','RC文氏桥：fo=1/(2piRC),|F|max=1/3,需Av=3。起振Av>3,稳幅Av=3。'),
('必记','起振信源=器件噪声+电源扰动。稳幅=NTC热敏电阻/JFET/二极管。'),
('区别','施密特(迟滞抗干扰)vs单门限(易抖动)。'),
],
}

tag_colors = {'易错': RGBColor(200,50,50), '必记': RGBColor(180,120,0), '公式': RGBColor(30,80,180), '区别': RGBColor(20,130,80)}

for ch, items in data.items():
    doc.add_heading(ch, level=1)
    for tag, text in items:
        p = doc.add_paragraph()
        r = p.add_run(f'[{tag}] ')
        r.bold = True; r.font.size = Pt(10)
        r.font.color.rgb = tag_colors.get(tag, RGBColor(0,0,0))
        p.add_run(text)

out = '模拟电子技术_关键注意事项.docx'
doc.save(out)
print(f'Done: {out}')
