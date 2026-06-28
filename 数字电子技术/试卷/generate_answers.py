#!/usr/bin/env python3
"""Generate answer documents for all 5 exam papers."""
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
import os, re

doc = Document()
style = doc.styles['Normal']
font = style.font
font.name = 'Microsoft YaHei'
font.size = Pt(11)
style.element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')

def add_title(text, level=0):
    h = doc.add_heading('', level=level)
    run = h.add_run(text)
    run.font.color.rgb = RGBColor(0x1a, 0x3a, 0x5c)
    return h

def add_q(question, answer, source=''):
    """Add a question-answer pair."""
    p = doc.add_paragraph()
    run = p.add_run('【原题】')
    run.bold = True
    run.font.color.rgb = RGBColor(0x2d, 0x7f, 0xc1)
    p.add_run(' ' + question)

    p2 = doc.add_paragraph()
    run2 = p2.add_run('【答案】')
    run2.bold = True
    run2.font.color.rgb = RGBColor(0xc0, 0x39, 0x2b)
    p2.add_run(' ' + answer)

    if source:
        p3 = doc.add_paragraph()
        run3 = p3.add_run('【来源】')
        run3.bold = True
        run3.font.color.rgb = RGBColor(0x27, 0xae, 0x60)
        p3.add_run(' ' + source)

    doc.add_paragraph('')

print("文档生成器已创建，5份文档将通过单独脚本生成")
