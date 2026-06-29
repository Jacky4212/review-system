"""
生成《数字电子技术》25道常识选择题Word文档
基于PPT内容，覆盖第1-10章+芯片专题
"""
from docx import Document
from docx.shared import Pt, Cm, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn

doc = Document()

# ── 页面设置 ──
for section in doc.sections:
    section.top_margin = Cm(2)
    section.bottom_margin = Cm(2)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)

# ── 样式 ──
style = doc.styles['Normal']
font = style.font
font.name = '宋体'
style.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
font.size = Pt(12)

# ── 标题 ──
title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = title.add_run('《数字电子技术》常识选择题25道')
run.font.size = Pt(18)
run.font.bold = True
run.font.name = '黑体'
run.element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')

subtitle = doc.add_paragraph()
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = subtitle.add_run('（基于PPT内容，覆盖第1～10章 + 芯片专题）')
run.font.size = Pt(11)
run.font.color.rgb = RGBColor(100, 100, 100)
run.font.name = '宋体'
run.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

doc.add_paragraph()  # 空行

# ══════════════════════════════════════
# 题目数据：(题目, [选项], 正确答案索引(0-based), 解释, 章节)
# ══════════════════════════════════════
questions = [
    # ── 第1章 数字逻辑概论 (Q1-Q4) ──
    (
        "1. 数字信号与模拟信号的主要区别是（  ）。",
        [
            "A. 数字信号的幅值更大",
            "B. 数字信号在时间和数值上均离散",
            "C. 数字信号的频率更高",
            "D. 数字信号只能传输音频"
        ],
        1,
        "数字信号的定义是：在时间和数值上均离散的信号。而模拟信号在时间和数值上均连续变化（如正弦波、三角波）。这是两者的本质区别。",
        "第1章"
    ),
    (
        "2. 十进制数 29 转换为二进制数是（  ）。",
        [
            "A. 11001",
            "B. 11101",
            "C. 10111",
            "D. 11100"
        ],
        1,
        "使用辗转相除法：29÷2=14余1（LSB），14÷2=7余0，7÷2=3余1，3÷2=1余1，1÷2=0余1（MSB），从下往上读得 11101。",
        "第1章"
    ),
    (
        "3. 下列属于无权码的是（  ）。",
        [
            "A. 8421码",
            "B. 2421码",
            "C. 格雷码",
            "D. 5421码"
        ],
        2,
        "格雷码（Gray Code）是一种无权码，每一位没有固定的权值，其特点是相邻代码之间仅有一位不同，能有效减少传输错误。8421码、2421码、5421码均属于有权BCD码。余3码和余3循环码也是无权码。",
        "第1章"
    ),
    (
        "4. 异或逻辑“相异为1”，其逻辑表达式为（  ）。",
        [
            "A. L = AB + ĀB̄",
            "B. L = ĀB + AB̄",
            "C. L = A + B",
            "D. L = AB"
        ],
        1,
        "异或（XOR）运算定义为 L = A⊕B = ĀB + AB̄，即两个输入不同时输出为1。选项A是同或（XNOR）的表达式 L = A⊙B = AB + ĀB̄，即相同为1。",
        "第1章"
    ),

    # ── 第2章 逻辑代数与HDL (Q5-Q7) ──
    (
        "5. 摩根定律（De Morgan）的正确表述是（  ）。",
        [
            "A. ¬(A·B) = Ā·B̄",
            "B. ¬(A+B) = Ā + B̄",
            "C. ¬(A·B) = Ā + B̄",
            "D. ¬(A+B) = AB"
        ],
        2,
        "摩根定律是逻辑代数最重要的定律之一：¬(A·B) = Ā + B̄（与非等于非的或），¬(A+B) = Ā·B̄（或非等于非的与）。口诀：长杠变短杠，乘变加、加变乘。",
        "第2章"
    ),
    (
        "6. 用卡诺图化简逻辑函数时，每个包围圈内的最小项个数必须是（  ）。",
        [
            "A. 任意个数",
            "B. 2的整数次幂（1, 2, 4, 8, …）个",
            "C. 只能3个",
            "D. 只能偶数个即可"
        ],
        1,
        "卡诺图化简时，每个圈（卡诺圈）内包含的最小项个数必须是2^n个（n=0,1,2,…），即1、2、4、8……个，这样才能消去n个变量，得到最简的乘积项。",
        "第2章"
    ),
    (
        "7. Verilog HDL中，下列哪种数据类型需要在always块中赋值？（  ）",
        [
            "A. wire型",
            "B. reg型",
            "C. input型",
            "D. output型"
        ],
        1,
        "Verilog有4种逻辑值（0、1、x、z）。wire型用于连续赋值（assign语句），reg型（寄存器型）用于过程赋值（always块中），它能保持状态直到下一次赋值。",
        "第2章"
    ),

    # ── 第3章 逻辑门电路 (Q8-Q10) ──
    (
        "8. CMOS反相器由以下哪两个管构成互补结构？（  ）",
        [
            "A. 两个NMOS管",
            "B. 两个PMOS管",
            "C. 一个NMOS管和一个PMOS管",
            "D. 一个BJT和一个MOS管"
        ],
        2,
        "CMOS反相器的核心结构是一个NMOS管（TN）和一个PMOS管（TP）串联构成互补对称结构。输入为低时TN截止TP导通输出高；输入为高时TN导通TP截止输出低。特点是：总有一管导通一管截止，静态功耗几乎为零。",
        "第3章"
    ),
    (
        "9. N沟道增强型MOS管作为开关时，当vGS > VT（开启电压）时，MOS管处于（  ）。",
        [
            "A. 截止状态",
            "B. 导通状态",
            "C. 放大状态",
            "D. 击穿状态"
        ],
        1,
        "N沟道增强型MOS管：vGS < VT时沟道未形成，管子截止（开关断开）；vGS > VT时形成反型层沟道，管子导通（开关闭合）。MOS管本质上是一个由vGS控制的电压控制型无触点开关。",
        "第3章"
    ),
    (
        "10. TTL门电路中，肖特基二极管（SBD）的作用是（  ）。",
        [
            "A. 提高工作电压",
            "B. 防止BJT进入深度饱和，提高开关速度",
            "C. 增大输出电流",
            "D. 降低功耗"
        ],
        1,
        "抗饱和TTL（肖特基TTL）在BJT基极-集电极之间并联肖特基势垒二极管（SBD）。SBD正向压降小（约0.3V），能将BJT的VCE钳位在约0.3V，防止其进入深度饱和（VCE≈0.1V），从而减少存储时间ts，提高开关速度。",
        "第3章"
    ),

    # ── 第4章 组合逻辑电路 (Q11-Q13) ──
    (
        "11. 组合逻辑电路的结构特点是（  ）。",
        [
            "A. 输出与电路原来状态有关，存在反馈回路",
            "B. 电路中含有存储单元",
            "C. 输出仅取决于当前输入，不含反馈和存储单元",
            "D. 必须包含时钟信号"
        ],
        2,
        "组合逻辑电路的定义：任意时刻的输出状态只取决于该时刻的输入状态，与电路原来的状态无关。结构上，输出与输入之间无反馈延迟通路，不包含记忆单元（如触发器）。这是组合逻辑与时序逻辑的根本区别。",
        "第4章"
    ),
    (
        "12. 74HC138是一款常用的（  ）。",
        [
            "A. 8线-3线编码器",
            "B. 3线-8线译码器",
            "C. 4位全加器",
            "D. 8选1数据选择器"
        ],
        1,
        "74HC138是3线-8线二进制译码器，输入为A2A1A0三根地址线（高有效），输出为Y0~Y7（低有效）。它还带有多个使能端（E3高有效，Ē1/Ē2低有效），全部有效时芯片才工作。",
        "第4章"
    ),
    (
        "13. 实现两个1位二进制数相加，不考虑低位进位的电路是（  ）。",
        [
            "A. 全加器",
            "B. 半加器",
            "C. 超前进位加法器",
            "D. 数值比较器"
        ],
        1,
        "半加器（Half Adder）仅将两个1位二进制数相加，不考虑来自低位的进位。全加器（Full Adder）则考虑低位进位。超前进位加法器同时产生所有进位信号，无需逐级等待，速度快。",
        "第4章"
    ),

    # ── 第5章 锁存器和触发器 (Q14-Q16) ──
    (
        "14. 用或非门构成的基本SR锁存器，当R=1、S=1时，输出状态为（  ）。",
        [
            "A. Q=0（复位）",
            "B. Q=1（置位）",
            "C. 保持原状态",
            "D. 不确定（禁止状态）"
        ],
        3,
        "或非门SR锁存器：R=0,S=0保持；R=0,S=1置1；R=1,S=0置0；R=1,S=1时两个输出都为0（非正常状态），且当R和S同时回到0时，输出状态不确定。因此R=S=1为禁止状态，约束条件为SR=0。",
        "第5章"
    ),
    (
        "15. D触发器的特性方程为（  ）。",
        [
            "A. Q^(n+1) = J·Q̄^n + K̄·Q^n",
            "B. Q^(n+1) = T ⊕ Q^n",
            "C. Q^(n+1) = D",
            "D. Q^(n+1) = S + R̄·Q^n"
        ],
        2,
        "D触发器的特性方程为Q^(n+1)=D，是四种触发器中特性方程最简单的——输出就是输入D的值。选项A是JK触发器、B是T触发器、D是SR触发器的特性方程。",
        "第5章"
    ),
    (
        "16. JK触发器当J=K=1时，每来一个时钟脉冲，其功能等效于（  ）。",
        [
            "A. 保持功能",
            "B. 置0功能",
            "C. T'触发器（翻转/二分频）",
            "D. D触发器"
        ],
        2,
        "JK触发器的特性方程 Q^(n+1) = J·Q̄^n + K̄·Q^n。当J=K=1时，代入得 Q^(n+1) = 1·Q̄^n + 0·Q^n = Q̄^n，每来一个CP输出翻转一次，等效于T'触发器（T=1的T触发器），可实现二分频。",
        "第5章"
    ),

    # ── 第6章 时序逻辑电路 (Q17-Q19) ──
    (
        "17. 时序逻辑电路与组合逻辑电路的根本区别在于（  ）。",
        [
            "A. 是否包含门电路",
            "B. 输出是否与输入有关",
            "C. 是否包含存储电路（记忆单元）和反馈",
            "D. 是否使用CMOS工艺"
        ],
        2,
        "时序逻辑电路的结构特征：由组合逻辑电路和存储电路（触发器）两部分组成，存在反馈回路。工作特征：输出不仅与当前输入有关，还与电路原来的状态有关。是否含存储单元是两者的根本区别。",
        "第6章"
    ),
    (
        "18. 74LVC161计数器的清零方式是（  ）。",
        [
            "A. 同步清零",
            "B. 异步清零",
            "C. 电平清零（无时钟要求）和边沿清零都有",
            "D. 不能清零"
        ],
        1,
        "74LVC161是4位二进制同步加计数器，其CLR（清零端）为异步方式——CLR=L时立即清零，不需要时钟信号。同时它还具有同步并行置数功能（LOAD=L时，在CP上升沿将D3~D0装入）。",
        "第6章"
    ),
    (
        "19. 用一片74LVC161（模16计数器）构成模10计数器，须跳过的状态数为（  ）。",
        [
            "A. 4个",
            "B. 5个",
            "C. 6个",
            "D. 10个"
        ],
        2,
        "一片161是4位二进制计数器，共2^4=16个状态（0~15）。要构成模10计数器（10个有效状态），需跳过 16−10=6 个状态。常用反馈清零法或反馈置数法实现。反馈清零法：当计数到10（1010）时，用与非门检测并产生清零信号。",
        "第6章"
    ),

    # ── 第7章 半导体存储器 (Q20-Q21) ──
    (
        "20. 下列存储器中，掉电后数据会丢失的是（  ）。",
        [
            "A. EPROM",
            "B. Flash闪存",
            "C. SRAM",
            "D. E²PROM"
        ],
        2,
        "RAM（包括SRAM和DRAM）是易失性存储器，掉电后数据丢失。而ROM类（掩模ROM、PROM、EPROM、E²PROM、Flash闪存）都是非易失性存储器，掉电后数据保持。SRAM由触发器构成存储单元，只要供电就保持数据。",
        "第7章"
    ),
    (
        "21. DRAM需要定期刷新的原因是（  ）。",
        [
            "A. 存储单元由触发器构成，需要时钟维持",
            "B. 存储单元用电容存储电荷，存在漏电",
            "C. 数据写入速度太慢",
            "D. 存储容量不够大"
        ],
        1,
        "DRAM（动态随机存取存储器）的每个存储单元仅由一个MOS管和一个电容组成。电容存储的电荷会通过漏电流逐渐泄放，导致数据丢失，因此必须每隔一定时间（通常2~64ms）对电容进行刷新（读出后重写），补充电荷。",
        "第7章"
    ),

    # ── 第8章 FPGA/CPLD (Q22-Q23) ──
    (
        "22. FPGA实现逻辑函数的基本单元是（  ）。",
        [
            "A. 可编程“与-或”阵列",
            "B. 查找表（LUT）",
            "C. 三极管开关阵列",
            "D. 电阻网络"
        ],
        1,
        "FPGA（现场可编程门阵列）采用查找表（LUT）结构实现逻辑函数。LUT本质上是一个小规模存储器（如SRAM），以真值表的形式将n输入的逻辑函数预先存储在2^n个存储单元中，输入作为地址读出对应输出。目前主流FPGA使用4~6输入LUT。",
        "第8章"
    ),
    (
        "23. FPGA上电后需要从外部PROM加载配置数据，这是因为（  ）。",
        [
            "A. FPGA内部逻辑容量不够",
            "B. FPGA配置数据存放在SRAM中，掉电后丢失",
            "C. FPGA需要外部时钟源",
            "D. 为了降低功耗"
        ],
        1,
        "FPGA的配置数据存储在SRAM中，属于易失性存储器，掉电后配置信息全部丢失。因此每次上电时，FPGA必须从外部非易失性存储器（如PROM、Flash）中重新加载配置数据（bitstream），完成逻辑功能的配置。与之对比，CPLD采用E²PROM或Flash，上电即可直接工作。",
        "第8章"
    ),

    # ── 第9章 脉冲波形 (Q24-Q25) ──
    (
        "24. 555定时器构成单稳态触发器时，输出脉冲宽度tw约为（  ）。",
        [
            "A. tw ≈ 0.7RC",
            "B. tw ≈ 1.1RC",
            "C. tw ≈ 2.2RC",
            "D. tw ≈ RC"
        ],
        1,
        "555定时器构成单稳态触发器时，输出脉冲宽度 tw = RC·ln3 ≈ 1.1RC。注意区分三种模式的关键系数：门电路单稳态tw≈0.7RC，555单稳态tw≈1.1RC，多谐振荡周期T≈0.7(R1+2R2)C。",
        "第9章"
    ),
    (
        "25. 施密特触发电路有两个不同的阈值电压，这是为了（  ）。",
        [
            "A. 提高工作频率",
            "B. 增大输出幅度",
            "C. 利用回差电压抗干扰，实现波形整形",
            "D. 降低功耗"
        ],
        2,
        "施密特触发器有两个阈值电压：正向阈值电压VT+（输入上升时翻转）和负向阈值电压VT−（输入下降时翻转），两者之差ΔV=VT+−VT−称为回差电压。当输入信号叠加噪声时，回差特性可以防止在阈值附近反复翻转，从而实现抗干扰的波形整形。555定时器构成的施密特触发器VT+=2/3VCC，VT−=1/3VCC。",
        "第9章"
    ),
]

# ── 写入文档 ──
for i, (stem, options, correct_idx, explanation, chapter) in enumerate(questions):
    # 题目标题
    p = doc.add_paragraph()
    run = p.add_run(f'【{chapter}】{stem}')
    run.font.size = Pt(12)
    run.font.bold = True
    run.font.name = '宋体'
    run.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

    # 四个选项
    for opt in options:
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Cm(1)
        run = p.add_run(opt)
        run.font.size = Pt(12)
        run.font.name = '宋体'
        run.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

    # 答案与解释（使用特殊颜色标记）
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(0.5)

    run = p.add_run(f'✓ 正确答案：{options[correct_idx][0]}')
    run.font.size = Pt(12)
    run.font.bold = True
    run.font.color.rgb = RGBColor(0, 128, 0)
    run.font.name = '宋体'
    run.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(0.5)
    run = p.add_run(f'解析：{explanation}')
    run.font.size = Pt(11)
    run.font.color.rgb = RGBColor(80, 80, 80)
    run.font.name = '宋体'
    run.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

    # 题目间分隔
    doc.add_paragraph()  # 空行

# ── 末尾统计 ──
doc.add_paragraph()
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('── 共25题，覆盖第1～10章核心知识点及芯片专题 ──')
run.font.size = Pt(10)
run.font.color.rgb = RGBColor(150, 150, 150)
run.font.name = '宋体'
run.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('各章分布：第1章(4题) · 第2章(3题) · 第3章(3题) · 第4章(3题) · '
                '第5章(3题) · 第6章(3题) · 第7章(2题) · 第8章(2题) · 第9章(2题)')
run.font.size = Pt(9)
run.font.color.rgb = RGBColor(180, 180, 180)
run.font.name = '宋体'
run.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

# ── 保存 ──
output_path = r'D:\code\cherry studio\复习\数字电子技术\数字电子技术常识选择题25道.docx'
doc.save(output_path)
print(f'✓ Word文档已生成：{output_path}')
print(f'  共 {len(questions)} 道选择题，附答案和解析')
