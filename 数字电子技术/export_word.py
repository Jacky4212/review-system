#!/usr/bin/env python3
"""Export ch07 and ch10 content to a Word document."""
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
import os, re

doc = Document()

style = doc.styles['Normal']
font = style.font
font.name = '微软雅黑'
font.size = Pt(11)
style.element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')

def add_title(text, level=0):
    h = doc.add_heading('', level=level)
    run = h.add_run(text)
    run.font.color.rgb = RGBColor(0x1a, 0x3a, 0x5c)
    return h

def add_bold_text(paragraph, bold_part, normal_part=''):
    run = paragraph.add_run(bold_part)
    run.bold = True
    if normal_part:
        paragraph.add_run(normal_part)

# ===== Title =====
add_title('数字电子技术基础 — 重点章节知识点', 0).alignment = WD_ALIGN_PARAGRAPH.CENTER
doc.add_paragraph('内容来源：华中科技大学《数字电子技术基础》课程PPT').alignment = WD_ALIGN_PARAGRAPH.CENTER
doc.add_paragraph('')

# ===== CH07 =====
add_title('第7章 半导体存储器', 1)

doc.add_heading('7.1 只读存储器（ROM）', level=2)
doc.add_paragraph('只读存储器（Read-Only Memory）在正常工作状态只能读出信息。断电后信息不会丢失，常用于存放固定信息（如程序、常数等）。')

doc.add_heading('基本概念', level=3)
p = doc.add_paragraph()
add_bold_text(p, '存储容量(M) = 字数 × 位数')
doc.add_paragraph('字：计算机中作为一个整体被存取传送处理的一组数据。')
doc.add_paragraph('地址：每个字的编号。')
doc.add_paragraph('字数 = 2ⁿ（n为地址线的总数）')
doc.add_paragraph('字长：一个字所含的位数。')

doc.add_heading('ROM的基本结构', level=3)
doc.add_paragraph('ROM主要由地址译码器、存储矩阵和输出控制电路三部分组成。')

doc.add_heading('ROM的分类', level=3)
doc.add_paragraph('• 固定ROM（掩模ROM）：出厂时数据已固化，不能更改。')
doc.add_paragraph('• PROM（一次可编程）：由带熔丝或反熔丝的器件构成，只能烧写一次。')
doc.add_paragraph('• EPROM（光可擦除）：用SIMOS管，紫外灯或X射线照15~20分钟擦除。')
doc.add_paragraph('• E²PROM（电可擦除）：在线电擦除，但集成度低。')
doc.add_paragraph('• Flash（闪存）：有NOR和NAND型两类。')

doc.add_heading('NOR Flash vs NAND Flash', level=3)
p = doc.add_paragraph()
add_bold_text(p, 'NOR型：', '列线上浮栅MOS漏极并联，源极接地的或非结构。按块擦除、按字写入，并行接口，常用于存放运行程序。')
p = doc.add_paragraph()
add_bold_text(p, 'NAND型：', '列线上浮栅MOS串联，集成度比NOR高。以块存取，串行接口，多用于固态硬盘、存储卡和U盘。')

doc.add_heading('7.2 随机存取存储器（RAM）', level=2)

doc.add_heading('SRAM（静态随机存取存储器）', level=3)
doc.add_paragraph('SRAM的存储单元是双稳态存储单元电路，由触发器构成。只要电源供电，数据就能保持。')

doc.add_heading('SSRAM（同步静态随机存取存储器）', level=3)
doc.add_paragraph('SSRAM是一种高速RAM。读写操作在时钟脉冲节拍控制下完成。')
doc.add_paragraph('• 普通模式读写（ADV=0）：按外部给定地址进行读/写操作。')
doc.add_paragraph('• 丛发模式读写（ADV=1）：由丛发计数器加1产生新地址。')

doc.add_heading('DRAM（动态随机存取存储器）', level=3)
doc.add_paragraph('DRAM的存储单元由一个MOS管T和一个电容器C组成。')
doc.add_paragraph('• 写入：T导通，若DI为1则向电容器充电（存1），反之放电（存0）。')
doc.add_paragraph('• 读出：T导通，C中数据通过位线和缓冲器输出。')
doc.add_paragraph('• 刷新：每次读出后必须及时刷新，因电容器上电荷会泄漏。')

doc.add_heading('FIFO与双口SRAM', level=3)
doc.add_paragraph('FIFO（First-in First-out）：先存入的数据先被读出，不能随机存取，用于高速数据采集缓冲器。')
doc.add_paragraph('双口SRAM：有两套完全独立完整的地址、数据和控制端口，共用同一个存储阵列。')

doc.add_heading('存储器容量的扩展', level=3)
doc.add_paragraph('• 字长扩展：利用芯片的并联方式实现。')
doc.add_paragraph('• 字数扩展：利用外加译码器控制芯片的片选输入端实现。')

# ===== CH10 =====
add_title('第10章 模数与数模转换器', 1)

doc.add_heading('10.1 D/A转换器', level=2)

doc.add_heading('D/A转换的基本思想', level=3)
doc.add_paragraph('二进制每位代码都有一定的权值，按其权的大小转换成模拟量，然后将这些模拟量相加，即可得到与数字量成正比的模拟量。')

doc.add_heading('倒T形电阻网络D/A转换器', level=3)
doc.add_paragraph('根据运放线性运用时虚地的概念，无论模拟开关Si处于何种位置，与Si相连的2R电阻将接"地"或虚地。')
doc.add_paragraph('流入每个2R电阻的电流从高位到低位按2的整数倍递减。')
doc.add_paragraph('输出电压与输入的二进制数成正比，实现了数字量到模拟量的转换。')

doc.add_heading('D/A转换器的主要技术指标', level=3)
doc.add_paragraph('• 分辨率：用位数表示。n位DAC最多有2ⁿ个模拟输出电压。')
doc.add_paragraph('• 转换精度：实际转换特性与理想转换特性之间的最大偏差。')
doc.add_paragraph('• 转换速度：用稳定时间和转换速率描述。')

doc.add_heading('10.2 A/D转换器', level=2)

doc.add_heading('A/D转换的工作过程', level=3)
doc.add_paragraph('A/D转换器一般要包括取样、保持、量化及编码4个过程。')

doc.add_heading('采样定理', level=3)
p = doc.add_paragraph()
add_bold_text(p, 'f_s ≥ 2f_imax')
doc.add_paragraph('采样信号S(t)的频率愈高，所采得信号经低通滤波器后愈能真实地复现输入信号。')

doc.add_heading('量化与量化误差', level=3)
doc.add_paragraph('数字信号在数值上是离散的。采样-保持电路的输出电压需按某种近似方式归化到离散电平上。')
p = doc.add_paragraph()
add_bold_text(p, '量化误差属原理误差，无法消除。')
doc.add_paragraph('A/D转换器的位数越多，各离散电平之间的差值越小，量化误差越小。')
doc.add_paragraph('• 只舍不入方式：最大量化误差为Δ')
doc.add_paragraph('• 四舍五入方式：最大量化误差为Δ/2')

doc.add_heading('三种A/D转换器比较', level=3)

doc.add_paragraph('① 并联比较型：转换速度最快（10ns~1μs），但电路复杂。')
doc.add_paragraph('② 逐次逼近型：转换速度适中（几μs~100μs），精度高，在速度和复杂度间达到很好的平衡。')
doc.add_paragraph('③ 双积分型：转换速度慢（几百μs~几ms），但抗干扰能力最强。')

doc.add_heading('A/D转换器的主要技术指标', level=3)
doc.add_paragraph('• 分辨率：以输出二进制数的位数表示。n位ADC能区分的最小电压为满量程输入的1/2ⁿ。')
doc.add_paragraph('• 转换精度：常用最低有效位(LSB)或满刻度的百分比(%FSR)表示。')
doc.add_paragraph('• 转换速度：用转换时间或转换速率描述。')

# ===== Save =====
base = os.path.dirname(os.path.abspath(__file__))
out_path = os.path.join(base, '第7章-第10章 知识点总结.docx')
doc.save(out_path)
print("文档已生成: " + out_path)
