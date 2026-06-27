#!/usr/bin/env python3
"""Generate a generalized reusable workflow document."""
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
import os

doc = Document()

# --- Style ---
style = doc.styles['Normal']
font = style.font
font.name = '微软雅黑'
font.size = Pt(11)
style.element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')

# ===== Title =====
title = doc.add_heading('', level=0)
run = title.add_run('PPT课程复习网站搭建 — 通用工作流')
run.font.size = Pt(20)
run.font.color.rgb = RGBColor(0x1a, 0x3a, 0x5c)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER

desc = doc.add_paragraph()
desc.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = desc.add_run('适用场景：从课程PPT中提取内容 → 搭建静态复习网站 → 部署到GitHub Pages')
run.font.color.rgb = RGBColor(0x7f, 0x8c, 0x8d)

doc.add_paragraph('')

# ===== 前置条件 =====
h = doc.add_heading('前置条件', level=1)
items = [
    'Python 3.8+ 环境（需安装 python-pptx 库）',
    'Git 已安装并配置好 GitHub 远程仓库',
    'GitHub CLI（gh）已安装并登录（可选，用于自动配置Pages）',
    'PPTX 格式的课件文件（.ppt 旧格式需先转换为 .pptx）',
]
for item in items:
    doc.add_paragraph(item, style='List Bullet')

doc.add_paragraph('')

# ===== 工作流 =====
doc.add_heading('通用工作流', level=1)

# ---------- Step 1 ----------
doc.add_heading('Step 1 — 提取 PPT 文字内容', level=2)

p = doc.add_paragraph()
p.add_run('目标：').bold = True
p.add_run('从所有 PPTX 文件中批量提取文字，保存为结构化文本文件。')

doc.add_paragraph('核心逻辑：', style='List Bullet')
core1 = [
    '遍历目标目录下所有 .pptx 文件',
    '对每个文件，逐页遍历所有形状（shapes）',
    '对每个形状：若有文本框（has_text_frame），提取段落文本；若有表格（has_table），逐行逐格提取',
    '按"来源文件 + 页码 + 文字"的格式写入 .txt 文件',
    '可设置过滤规则：如跳过重复文件、(2)副本等',
]
for c in core1:
    doc.add_paragraph(c, style='List Number')

p = doc.add_paragraph()
p.add_run('关键库：').bold = True
p.add_run('python-pptx (from pptx import Presentation)')

p = doc.add_paragraph()
p.add_run('伪代码：').bold = True
code = """for each .pptx file:
    prs = Presentation(file)
    for each slide in prs.slides:
        for each shape in slide.shapes:
            if shape.has_text_frame:
                extract paragraphs
            if shape.has_table:
                extract cells row by row
        save to {file_name}.txt"""
doc.add_paragraph(code, style='No Spacing')

doc.add_paragraph('')
p = doc.add_paragraph()
p.add_run('输出结果：').bold = True
p.add_run('每个 PPTX 对应一个 .txt 文件，保存到指定输出目录。')

# ---------- Step 2 ----------
doc.add_heading('Step 2 — 录音转写（可选）', level=2)

p = doc.add_paragraph()
p.add_run('目标：').bold = True
p.add_run('将课程录音（.m4a/.mp3/.wav）转写为文字，作为复习补充材料。')

doc.add_paragraph('方案一（推荐）：', style='List Bullet')
steps2a = [
    '安装 openai-whisper 或 faster-whisper：pip install faster-whisper',
    '若网络受限，单独下载模型文件：用 curl -k 从 HuggingFace Hub 手动下载模型到本地路径',
    '模型文件包括：config.json, model.bin, tokenizer.json, vocabulary.txt',
    '加载本地模型：WhisperModel(model_path, device="cpu", compute_type="int8", local_files_only=True)',
    '调用 model.transcribe(audio_file, language="zh") 进行中文转写',
    '遍历 segments，拼接全文，保存为 .txt 文件',
]
for s in steps2a:
    doc.add_paragraph(s, style='List Number')

doc.add_paragraph('方案二（轻量离线）：', style='List Bullet')
steps2b = [
    '安装 vosk：pip install vosk',
    '从 alphacephei.com 下载中文模型（vosk-model-small-cn-0.22，约42MB）',
    '加载模型，对音频文件进行流式或一次性识别',
]
for s in steps2b:
    doc.add_paragraph(s, style='List Number')

p = doc.add_paragraph()
p.add_run('注意：').bold = True
p.add_run('中文转写建议指定 language="zh" 以提高准确率。首次运行需下载模型文件（~75MB）。')

# ---------- Step 3 ----------
doc.add_heading('Step 3 — 构建静态复习网站', level=2)

p = doc.add_paragraph()
p.add_run('目标：').bold = True
p.add_run('基于提取的文字内容，构建可直接在浏览器中打开的知识梳理网站。')

doc.add_paragraph('设计原则：', style='List Bullet')
principles = [
    '纯静态 HTML + CSS，零外部依赖，无需服务器',
    '内容忠实于 PPT 原话，标注来源',
    '响应式设计，兼容桌面和移动端',
    '按章节组织，导航清晰，可扩展',
]
for p_text in principles:
    doc.add_paragraph(p_text, style='List Number')

p = doc.add_paragraph()
run = p.add_run('\n⚠️ 核心原则（重中之重）：')
run.bold = True
run.font.color.rgb = RGBColor(0xc0, 0x39, 0x2b)
p2 = doc.add_paragraph()
run2 = p2.add_run('每个章节的知识点都必须人工理解PPT原文后手动精写，绝不能用脚本批量生成。')
run2.bold = True
steps_warn = [
    '通读该章所有PPT提取文本，理解知识结构和逻辑关系',
    '筛选出核心概念、定义、公式、功能表，舍弃无关细节',
    '用自己的语言组织，但关键术语和定义必须贴近PPT原话',
    '判断哪些内容需要上下标、表格、高亮框等特殊格式',
    '手动编写HTML，每个章节单独打磨',
]
for s in steps_warn:
    doc.add_paragraph(s, style='List Number')

doc.add_paragraph('推荐文件结构：', style='List Bullet')
structure = """project-root/
  ├── index.html          ← 课程首页（导航卡片）
  ├── ch03.html           ← 各章节页面（按需创建）
  ├── ch07.html
  ├── ch10.html
  └── chips.html          ← 专题页面（如常用芯片/公式速查）"""
doc.add_paragraph(structure, style='No Spacing')

doc.add_paragraph('')
doc.add_paragraph('页面内容组织规范：', style='List Bullet')
content_rules = [
    '每页顶部：教学基本要求（直接从PPT摘录）',
    '正文：按 PPT 小节划分，使用 h2/h3 层级标题',
    '关键概念用 def-box 高亮框展示定义',
    '功能表/真值表用 HTML table 呈现',
    '公式用居中 formula 块展示',
    '多芯片或方法对比用卡片网格（compare-grid）布局',
    '页脚注明内容来源章节',
    '章节间提供上下页导航链接',
]
for r in content_rules:
    doc.add_paragraph(r, style='List Number')

doc.add_paragraph('')
p = doc.add_paragraph()
p.add_run('HTML重建规范（关键）：').bold = True

doc.add_paragraph('')
p = doc.add_paragraph()
p.add_run('1. 表格重建：').bold = True
p.add_run('PPT提取出的表格通常是管道分隔的纯文本：')
doc.add_paragraph('清零 | 预置 | 使能 | 时钟 | 输出', style='No Spacing')
doc.add_paragraph('L   | ×   | ×   | ×   | L L L L', style='No Spacing')
doc.add_paragraph('')
p = doc.add_paragraph()
p.add_run('处理方法：').bold = True
doc.add_paragraph('识别表头行和数�的行，手动重建为 <table>')
doc.add_paragraph('表头用 <th>，数据行用 <td>，对齐用 text-align: center')
doc.add_paragraph('功能表逻辑：H=高电平 L=低电平 ×=任意 ↑=上升沿')
doc.add_paragraph('')

p = doc.add_paragraph()
p.add_run('2. 上下标处理：').bold = True
p.add_run('PPT文本中的VGS、2n等写法，需根据上下文还原为正确的上下标：')
doc.add_paragraph('文本"VGS" → HTML: v<sub>GS</sub> （GS是下标）', style='List Bullet')
doc.add_paragraph('文本"2n"  → HTML: 2<sup>n</sup> （n是上标）', style='List Bullet')
doc.add_paragraph('文本"VREF" → HTML: V<sub>REF</sub> （REF是下标）', style='List Bullet')
doc.add_paragraph('文本"VDD" → HTML: V<sub>DD</sub>', style='List Bullet')
doc.add_paragraph('文本"Qn+1" → HTML: Q<sub>n+1</sub>', style='List Bullet')
doc.add_paragraph('文本"VOH(min)" → HTML: V<sub>OH</sub>(min)', style='List Bullet')
doc.add_paragraph('文本"VIL(max)" → HTML: V<sub>IL</sub>(max)', style='List Bullet')
doc.add_paragraph('')
p = doc.add_paragraph()
p.add_run('判断依据：').bold = True
p.add_run('大写字母组合跟在V/I/Q等后面 → 下标；数字/变量跟在幂次后面 → 上标；括号内容如(min)/(max) → 正常文本。')

doc.add_paragraph('')
p = doc.add_paragraph()
p.add_run('3. 公式呈现：').bold = True
p.add_run('提取文本中的公式如 tw=RC1n3≈1.1RC，手动格式化为：')
doc.add_paragraph('使用 div.formula 居中展示', style='List Bullet')
doc.add_paragraph('变量用 <i> 斜体或直接文本', style='List Bullet')
doc.add_paragraph('约等号≈、乘号×、箭头→等保留原符号', style='List Bullet')
doc.add_paragraph('关键公式用 .highlight 黄色高亮框突出', style='List Bullet')
doc.add_paragraph('')

p = doc.add_paragraph()
p.add_run('4. CSS类名规范：').bold = True
doc.add_paragraph('.def-box — 定义/概念高亮框（浅灰背景）', style='List Bullet')
doc.add_paragraph('.highlight — 重要提示（黄色背景+左侧黄边）', style='List Bullet')
doc.add_paragraph('.formula — 居中公式展示', style='List Bullet')
doc.add_paragraph('.tag-xxx — 芯片/技术标签（如 .tag-rom .tag-adc）', style='List Bullet')
doc.add_paragraph('.compare-grid + .compare-card — 对比卡片网格', style='List Bullet')
doc.add_paragraph('table + th/td — 功能表/真值表', style='List Bullet')

# ---------- Step 4 ----------
doc.add_heading('Step 4 — 版本控制与推送到 GitHub', level=2)

p = doc.add_paragraph()
p.add_run('目标：').bold = True
p.add_run('将网站文件纳入 Git 管理并推送到远程仓库。')

steps4 = [
    'cd 到仓库根目录',
    'git add <网站文件> 添加新文件',
    'git commit -m "feat: 课程复习网站" 提交',
    'git push origin <branch> 推送到 GitHub',
    '若仓库根目录已有 index.html，建议改为导航页，链接到各课程站点',
]
for s in steps4:
    doc.add_paragraph(s, style='List Number')

p = doc.add_paragraph()
p.add_run('建议：').bold = True
p.add_run('仓库根目录下的 index.html 可作为多课程统一入口（导航页），每个课程一个子目录。')

# ---------- Step 5 ----------
doc.add_heading('Step 5 — 启用 GitHub Pages', level=2)

p = doc.add_paragraph()
p.add_run('目标：').bold = True
p.add_run('开启 GitHub Pages，使网站可通过公网域名直接访问。')

steps5 = [
    '方式一（gh CLI）：gh api repos/{owner}/{repo}/pages --method POST --field source=\'{"branch":"master","path":"/"}\'',
    '方式二（网页端）：进入仓库 Settings → Pages → 选择分支和目录 → Save',
    '等待构建完成（状态变为 "built"），通常需 1~5 分钟',
    '访问 https://{owner}.github.io/{repo}/ 即可看到站点',
]
for s in steps5:
    doc.add_paragraph(s, style='List Number')

p = doc.add_paragraph()
p.add_run('注意：').bold = True
p.add_run('GitHub Pages 默认使用 Jekyll，但纯静态 HTML 网站可直接工作。如果路径中带下划线，需在根目录添加 .nojekyll 文件。')

# ---------- Step 6 ----------
doc.add_heading('Step 6 — 维护与扩展', level=2)

p = doc.add_paragraph()
p.add_run('目标：').bold = True
p.add_run('持续更新网站内容，适配更多课程。')

steps6 = [
    '增加新章节：参照现有 HTML 模板创建新文件，在 index.html 导航栏添加卡片',
    '增加新专题：创建独立的专题页面（如 chips.html），引用各章节相关内容',
    '更新内容：重新运行提取脚本获得最新PPT内容，更新对应章节页面',
    'git commit & push 后，GitHub Pages 自动重新部署',
]
for s in steps6:
    doc.add_paragraph(s, style='List Number')

# ===== 关键技巧 =====
doc.add_heading('常见问题与技巧', level=1)

tips = [
    ('SSL证书问题', '在受限网络环境中，Python 的 httpx/requests 可能因 SSL 验证失败无法连接 HuggingFace。解决方案：使用 curl -k 手动下载模型文件，或设置环境变量 CURL_CA_BUNDLE="" 后重试。'),
    ('中文乱码', 'Windows 终端输出中文可能乱码。解决方案：将输出写入 UTF-8 编码的文件，用 Read 工具或直接打开文件查看。'),
    ('GitHub Pages 不更新', '推送后等待 1~5 分钟。若长时间未更新，可尝试在 Pages 设置页面触发重新部署。'),
    ('大文件推送', '模型文件（whisper-tiny-model ~75MB, vosk ~42MB）不推荐推送到 Git 仓库。建议在 .gitignore 中排除，或使用 Git LFS。'),
    ('PPT表格提取不完整', '部分复杂表格（合并单元格、嵌套表格）可能导致提取结果不完整。建议手动复核关键功能表。'),
]
for title_text, detail in tips:
    p = doc.add_paragraph()
    p.add_run(f'【{title_text}】').bold = True
    p.add_run(f' {detail}')

# ===== 工具清单 =====
doc.add_heading('工具与依赖清单', level=1)

doc.add_paragraph('Python 库：', style='List Bullet')
libs = [
    'python-pptx — PPTX 文件文字提取',
    'faster-whisper — 语音转写（GPU/CPU），依赖 ctranslate2, huggingface-hub',
    'vosk — 轻量离线语音识别',
    'python-docx — Word 文档生成',
]
for lib in libs:
    doc.add_paragraph(lib, style='List Number')

doc.add_paragraph('命令行工具：', style='List Bullet')
cmds = [
    'curl — 手动下载模型文件',
    'gh — GitHub CLI，用于操作仓库和 Pages',
    'git — 版本控制',
]
for cmd in cmds:
    doc.add_paragraph(cmd, style='List Number')

# ===== Footer =====
doc.add_paragraph('')
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('— 本文档描述通用流程，具体路径和参数请根据实际环境调整 —')
run.font.color.rgb = RGBColor(0x95, 0xa5, 0xa6)

# ===== Save =====
out_path = os.path.join(os.path.dirname(os.path.abspath(__file__)), 'PPT课程复习网站搭建-通用工作流.docx')
doc.save(out_path)
print("文档已生成: " + out_path)
