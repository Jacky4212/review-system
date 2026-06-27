#!/usr/bin/env python3
"""Export truth tables and key info for 555, 161, 138 chips to Word."""
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
import os

doc = Document()

style = doc.styles['Normal']
font = style.font
font.name = 'Microsoft YaHei'
font.size = Pt(11)
style.element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')

def add_title(text, level=0):
    h = doc.add_heading('', level=level)
    run = h.add_run(text)
    run.font.color.rgb = RGBColor(0x1a, 0x3a, 0x5c)
    return h

def add_tab(rows_data):
    """Add a table from list of lists."""
    rows = len(rows_data)
    cols = max(len(r) for r in rows_data)
    table = doc.add_table(rows=rows, cols=cols, style='Light Grid Accent 1')
    for i, row_data in enumerate(rows_data):
        for j, cell_text in enumerate(row_data):
            table.rows[i].cells[j].text = str(cell_text)

# ===== Title =====
add_title('数字电子技术 — 常用芯片真值表与功能表', 0).alignment = WD_ALIGN_PARAGRAPH.CENTER
doc.add_paragraph('内容来源：华中科技大学《数字电子技术基础》课程PPT').alignment = WD_ALIGN_PARAGRAPH.CENTER
doc.add_paragraph('')

# ===== 555 =====
add_title('555 定时器', 1)
doc.add_paragraph('555定时器是一种应用方便的中规模集成电路，广泛用于信号的产生、变换、控制与检测。有双极性和CMOS之分。')

doc.add_heading('内部结构', level=2)
doc.add_paragraph('555定时器内部由五部分组成：')
doc.add_paragraph('1. 电阻分压器——三个等值电阻串联，提供参考电压（2/3 Vcc 和 1/3 Vcc）')
doc.add_paragraph('2. 电压比较器——两个比较器C1和C2')
doc.add_paragraph('3. 基本SR锁存器')
doc.add_paragraph('4. 输出缓冲反相器')
doc.add_paragraph('5. 集电极开路输出三极管')

doc.add_heading('功能表', level=2)
add_tab([
    ['阈值输入 (TH)', '触发输入 (TR)', '复位 (R)', '输出 (OUT)', '放电管 (DIS)'],
    ['×', '×', 'L', 'L', '导通'],
    ['> 2/3 Vcc', '> 1/3 Vcc', 'H', 'L', '导通'],
    ['< 2/3 Vcc', '> 1/3 Vcc', 'H', '不变', '不变'],
    ['< 2/3 Vcc', '< 1/3 Vcc', 'H', 'H', '截止'],
])

doc.add_heading('三种典型应用', level=2)
doc.add_paragraph('1. 施密特触发电路')
doc.add_paragraph('   正向阈值电压 VT+ = 2/3 Vcc')
doc.add_paragraph('   负向阈值电压 VT- = 1/3 Vcc')
doc.add_paragraph('   回差电压 Delta_VT = 1/3 Vcc')

doc.add_paragraph('2. 单稳态电路')
doc.add_paragraph('   稳态：输出为0')
doc.add_paragraph('   触发：TR < 1/3 Vcc时翻转到暂稳态，输出为1')
doc.add_paragraph('   暂稳态持续时间：tw = RC * ln3 = 1.1 RC')

doc.add_paragraph('3. 多谐振荡电路')
doc.add_paragraph('   充电时间（高电平）：tpH = 0.7 (R1 + R2) C')
doc.add_paragraph('   放电时间（低电平）：tpL = 0.7 R2 C')
doc.add_paragraph('   振荡周期：T = 0.7 (R1 + 2R2) C')

# ===== 161 =====
add_title('74LVC161 计数器', 1)
doc.add_paragraph('74LVC161是4位二进制同步加计数器，具有异步清零和同步并行置数功能。')

doc.add_heading('芯片特点', level=2)
doc.add_paragraph('• 异步清零：清零端(CLR)为低电平时，立即清零。')
doc.add_paragraph('• 同步并行置数：预置端(LOAD)为低电平时，在CP上升沿将D3~D0数据装入。')
doc.add_paragraph('• 计数使能：CET和CEP同时为高电平时，CP上升沿加1计数。')
doc.add_paragraph('• 并行进位输出：TC = CET * Q3 * Q2 * Q1 * Q0')

doc.add_heading('逻辑功能表', level=2)
add_tab([
    ['清零', '预置', 'CEP', 'CET', '时钟CP', 'D3D2D1D0', 'Q3Q2Q1Q0', 'TC'],
    ['L', 'x', 'x', 'x', 'x', 'x', 'L L L L', 'L'],
    ['H', 'L', 'x', 'x', 'up-arrow', 'D3~D0', 'D3D2D1D0', '*'],
    ['H', 'H', 'L', 'x', 'x', 'x', '保持', '*'],
    ['H', 'H', 'x', 'L', 'x', 'x', '保持', '*'],
    ['H', 'H', 'H', 'H', 'up-arrow', 'x', '计数(加1)', '*'],
])

doc.add_heading('构成任意进制计数器的方法', level=2)
doc.add_paragraph('反馈清零法：利用异步清零输入端，在计数过程中跳过M-N个状态。')
doc.add_paragraph('反馈置数法：利用同步置数端，在计数过程中跳过M-N个状态。')
doc.add_paragraph('多片级联：并行进位（低位进位作高位使能）或串行进位（低位进位作高位时钟）。')

# ===== 138 =====
add_title('74HC138 译码器', 1)
doc.add_paragraph('74HC138是3线-8线二进制译码器。译码是编码的逆过程，能将二进制码翻译成代表某一特定含义的信号。')

doc.add_heading('芯片特点', level=2)
doc.add_paragraph('• 3个输入端 A2A1A0（高电平有效）')
doc.add_paragraph('• 8个输出端 Y0~Y7（低电平有效）')
doc.add_paragraph('• 3个使能端：E1、E2低有效，E3高有效')
doc.add_paragraph('• 输出低有效：被选中的输出为L，其余为H')

doc.add_heading('功能表', level=2)
add_tab([
    ['E3', 'E1', 'E2', 'A2', 'A1', 'A0', 'Y0', 'Y1', 'Y2', 'Y3', 'Y4', 'Y5', 'Y6', 'Y7'],
    ['x', 'H', 'x', 'x', 'x', 'x', 'H', 'H', 'H', 'H', 'H', 'H', 'H', 'H'],
    ['x', 'x', 'H', 'x', 'x', 'x', 'H', 'H', 'H', 'H', 'H', 'H', 'H', 'H'],
    ['L', 'x', 'x', 'x', 'x', 'x', 'H', 'H', 'H', 'H', 'H', 'H', 'H', 'H'],
    ['H', 'L', 'L', 'L', 'L', 'L', 'L', 'H', 'H', 'H', 'H', 'H', 'H', 'H'],
    ['H', 'L', 'L', 'L', 'L', 'H', 'H', 'L', 'H', 'H', 'H', 'H', 'H', 'H'],
    ['H', 'L', 'L', 'L', 'H', 'L', 'H', 'H', 'L', 'H', 'H', 'H', 'H', 'H'],
    ['H', 'L', 'L', 'L', 'H', 'H', 'H', 'H', 'H', 'L', 'H', 'H', 'H', 'H'],
    ['H', 'L', 'L', 'H', 'L', 'L', 'H', 'H', 'H', 'H', 'L', 'H', 'H', 'H'],
    ['H', 'L', 'L', 'H', 'L', 'H', 'H', 'H', 'H', 'H', 'H', 'L', 'H', 'H'],
    ['H', 'L', 'L', 'H', 'H', 'L', 'H', 'H', 'H', 'H', 'H', 'H', 'L', 'H'],
    ['H', 'L', 'L', 'H', 'H', 'H', 'H', 'H', 'H', 'H', 'H', 'H', 'H', 'L'],
])

doc.add_paragraph('注：H=高电平，L=低电平，x=任意。三个使能端全部有效(E3=H, E1=L, E2=L)时芯片工作。')

doc.add_heading('译码器的扩展', level=2)
doc.add_paragraph('• 用2片74X138可以构成4线-16线译码器')
doc.add_paragraph('• 用74X139和74X138可以构成5线-32线译码器')

doc.add_heading('应用', level=2)
doc.add_paragraph('• 地址译码：常用于识别不同设备，每个输出对应一个设备的片选信号。')
doc.add_paragraph('• 数据分配器：将公共数据线上的数据按需要送到不同的通道。')

# ===== Summary =====
add_title('三芯片快速对比', 1)
add_tab([
    ['芯片', '功能', '输入', '输出', '关键特点'],
    ['555', '定时器/多谐振荡', 'TH, TR, R, CV', 'OUT, DIS', '三种工作模式'],
    ['74LVC161', '4位二进制计数器', 'CP, CLR, LOAD, CET, CEP, D3~D0', 'Q3~Q0, TC', '异步清零+同步置数'],
    ['74HC138', '3线-8线译码器', 'A2A1A0, E1E2E3', 'Y0~Y7', '低有效输出+多使能'],
])

# ===== Save =====
out_path = os.path.join(os.path.dirname(os.path.abspath(__file__)), '常用芯片真值表-555_161_138.docx')
doc.save(out_path)
print("文档已生成: " + out_path)
